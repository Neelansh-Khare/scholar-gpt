import io
from typing import List, Dict, Tuple
import fitz
from langchain_core.documents import Document

def extract_text_from_pdf(pdf_bytes: bytes, filename: str) -> Tuple[List[Document], Dict[str, any]]:
    documents = []
    metadata = {
        "filename": filename,
        "total_pages": 0,
        "total_text_length": 0,
        "format": "PDF"
    }
    
    try:
        pdf_document = fitz.open(stream=pdf_bytes, filetype="pdf")
        metadata["total_pages"] = pdf_document.page_count
        
        for page_num in range(pdf_document.page_count):
            page = pdf_document[page_num]
            text = page.get_text()
            
            text = clean_text(text)
            
            if text.strip():
                doc = Document(
                    page_content=text,
                    metadata={
                        "source": filename,
                        "page": page_num + 1,
                        "total_pages": pdf_document.page_count,
                        "format": "PDF"
                    }
                )
                documents.append(doc)
                metadata["total_text_length"] += len(text)
        
        pdf_document.close()
        
    except Exception as e:
        raise Exception(f"Error processing PDF {filename}: {str(e)}")
    
    return documents, metadata

def extract_text_from_txt(txt_bytes: bytes, filename: str) -> Tuple[List[Document], Dict[str, any]]:
    metadata = {
        "filename": filename,
        "total_pages": 1,
        "total_text_length": 0,
        "format": "TXT"
    }
    
    try:
        text = txt_bytes.decode('utf-8', errors='ignore')
        text = clean_text(text)
        
        if not text.strip():
            return [], metadata
        
        doc = Document(
            page_content=text,
            metadata={
                "source": filename,
                "page": 1,
                "total_pages": 1,
                "format": "TXT"
            }
        )
        metadata["total_text_length"] = len(text)
        return [doc], metadata
        
    except Exception as e:
        raise Exception(f"Error processing TXT {filename}: {str(e)}")

def extract_text_from_docx(docx_bytes: bytes, filename: str) -> Tuple[List[Document], Dict[str, any]]:
    try:
        from docx import Document as DocxDocument
    except ImportError:
        raise Exception("python-docx not installed. Install with: pip install python-docx")
    
    metadata = {
        "filename": filename,
        "total_pages": 1,
        "total_text_length": 0,
        "format": "DOCX"
    }
    
    try:
        doc_file = io.BytesIO(docx_bytes)
        doc = DocxDocument(doc_file)
        
        full_text = "\n".join([para.text for para in doc.paragraphs])
        full_text = clean_text(full_text)
        
        if not full_text.strip():
            return [], metadata
        
        document = Document(
            page_content=full_text,
            metadata={
                "source": filename,
                "page": 1,
                "total_pages": 1,
                "format": "DOCX"
            }
        )
        metadata["total_text_length"] = len(full_text)
        return [document], metadata
        
    except Exception as e:
        raise Exception(f"Error processing DOCX {filename}: {str(e)}")

def extract_text_from_html(html_bytes: bytes, filename: str) -> Tuple[List[Document], Dict[str, any]]:
    try:
        from html.parser import HTMLParser
        import re
    except ImportError:
        raise Exception("HTML parsing library not available")
    
    metadata = {
        "filename": filename,
        "total_pages": 1,
        "total_text_length": 0,
        "format": "HTML"
    }
    
    try:
        html_text = html_bytes.decode('utf-8', errors='ignore')
        text = re.sub('<[^<]+?>', '', html_text)
        text = clean_text(text)
        
        if not text.strip():
            return [], metadata
        
        document = Document(
            page_content=text,
            metadata={
                "source": filename,
                "page": 1,
                "total_pages": 1,
                "format": "HTML"
            }
        )
        metadata["total_text_length"] = len(text)
        return [document], metadata
        
    except Exception as e:
        raise Exception(f"Error processing HTML {filename}: {str(e)}")

def clean_text(text: str) -> str:
    text = text.replace('\x00', '')
    
    lines = text.split('\n')
    cleaned_lines = []
    for line in lines:
        line = ' '.join(line.split())
        if line:
            cleaned_lines.append(line)
    
    text = '\n'.join(cleaned_lines)
    
    text = text.strip()
    
    return text

def process_multiple_files(uploaded_files) -> Tuple[List[Document], List[Dict[str, any]]]:
    all_documents = []
    all_metadata = []
    
    for uploaded_file in uploaded_files:
        try:
            file_bytes = uploaded_file.read()
            file_name = uploaded_file.name.lower()
            
            if file_name.endswith('.pdf'):
                documents, metadata = extract_text_from_pdf(file_bytes, uploaded_file.name)
            elif file_name.endswith('.txt'):
                documents, metadata = extract_text_from_txt(file_bytes, uploaded_file.name)
            elif file_name.endswith('.docx'):
                documents, metadata = extract_text_from_docx(file_bytes, uploaded_file.name)
            elif file_name.endswith(('.html', '.htm')):
                documents, metadata = extract_text_from_html(file_bytes, uploaded_file.name)
            else:
                raise Exception(f"Unsupported file format: {file_name}")
            
            all_documents.extend(documents)
            all_metadata.append(metadata)
            
        except Exception as e:
            all_metadata.append({
                "filename": uploaded_file.name,
                "error": str(e),
                "total_pages": 0,
                "total_text_length": 0,
                "format": "Unknown"
            })
    
    return all_documents, all_metadata

def process_multiple_pdfs(uploaded_files) -> Tuple[List[Document], List[Dict[str, any]]]:
    return process_multiple_files(uploaded_files)
